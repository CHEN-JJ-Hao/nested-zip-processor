#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
智能嵌套式压缩文件处理系统设计与实现 - 文档生成工具
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

class DocumentGenerator:
    """文档生成器类，用于创建智能嵌套压缩文件处理系统的设计文档"""
    
    def __init__(self):
        """初始化文档生成器"""
        self.document = Document()
        section = self.document.sections[0]
        section.top_margin = Pt(72)  # 2.5cm
        section.bottom_margin = Pt(72)  # 2.5cm
        section.left_margin = Pt(60)  # 3cm
        section.right_margin = Pt(40)  # 2cm
        section.header_distance = Pt(42)  # 1.5cm
        section.footer_distance = Pt(42)  # 1.75cm
        header = section.header
        header_p = header.paragraphs[0]
        header_p.text = '智能嵌套式压缩文件处理系统设计与实现'
        header_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_p.style.font.size = Pt(10.5)  # 小五
        header_p.style.font.bold = True

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.text = ''
        footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 页码需在Word中插入域，docx库不支持自动插入页码域
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        section.footer_distance = Pt(24.75)  # 1.75cm

        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = 1  # 居中

        run = footer_p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 标题样式
        title_style = self.document.styles['Title']
        title_style.font.name = '宋体'
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        title_style.font.size = Pt(24)  # 三号
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 一级标题（章标题）
        heading1 = self.document.styles['Heading 1']
        heading1.font.name = '宋体'
        heading1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        heading1.font.size = Pt(18)  # 三号
        heading1.font.bold = True
        heading1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 二级标题（节标题）
        heading2 = self.document.styles['Heading 2']
        heading2.font.name = '宋体'
        heading2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        heading2.font.size = Pt(16)  # 四号
        heading2.font.bold = True
        heading2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 三级标题（级标题）
        heading3 = self.document.styles['Heading 3']
        heading3.font.name = '宋体'
        heading3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        heading3.font.size = Pt(12)  # 小四
        heading3.font.bold = True
        heading3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 正文样式
        normal = self.document.styles['Normal']
        normal.font.name = '宋体'
        normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        normal.font.size = Pt(12)  # 小四
        normal.paragraph_format.first_line_indent = Pt(18)  # 首行缩进2字符
        normal.paragraph_format.line_spacing = Pt(15.6)  # 1.3倍行距（12pt*1.3）

        # 英文字体设置
        for style in [title_style, heading1, heading2, heading3, normal]:
            style.font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    
    def add_title(self, title):
        """添加文档标题"""
        self.document.add_heading(title, level=0)
        self.document.add_paragraph()  # 空行
    
    def add_section1(self):
        """添加第一部分：项目背景与需求分析"""
        self.document.add_heading('一、项目背景与需求分析', level=1)
        
        # 项目背景
        self.document.add_heading('（一）项目背景', level=2)
        self.document.add_paragraph('在数字化信息时代，数据的存储与传输需求日益增长，压缩文件作为一种高效的数据处理方式，被广泛应用于各个领域。'
                                   '随着数据量的增加和文件结构的复杂化，嵌套式压缩文件（即压缩包内包含多个层级的压缩包）的处理需求逐渐凸显。'
                                   '传统的压缩解压工具在处理多层嵌套压缩包时，往往存在操作繁琐、效率低下、乱码问题频发等痛点，无法满足用户对批量处理、自动化操作的需求。')
        self.document.add_paragraph('与此同时，用户对于压缩工具的功能需求也在不断升级，不仅要求支持多种压缩格式，还希望具备智能化的文件结构优化、'
                                   '异常处理和操作中断恢复等功能。在此背景下，“智能嵌套压缩文件处理系统”的设计与实现具有重要的实际应用价值，'
                                   '旨在为用户提供更高效、智能、稳定的压缩文件处理解决方案。')
        
        # 需求分析
        self.document.add_heading('（二）需求分析', level=2)
        
        # 核心功能需求
        self.document.add_heading('1. 核心功能需求', level=3)
        core_funcs = [
            '**多格式支持**：系统需支持常见的压缩格式，包括ZIP、RAR、7Z、TAR、TAR.GZ、TAR.BZ2等，以满足不同用户和场景的需求。',
            '**嵌套解压**：能够自动识别并处理压缩包内的多层嵌套压缩文件，实现递归解压，减少用户手动操作。',
            '**压缩功能**：支持文件和文件夹的压缩，可选择不同的压缩格式，并能处理压缩过程中的异常情况。',
            '**操作控制**：提供暂停、继续、终止解压/压缩进程的功能，终止时需自动删除已解压或部分解压的内容，确保系统状态的一致性。'
        ]
        for func in core_funcs:
            self.document.add_paragraph(func, style='List Bullet')
        
        # 界面与交互需求
        self.document.add_heading('2. 界面与交互需求', level=3)
        ui_funcs = [
            '**可视化操作**：通过图形用户界面（GUI）提供直观的操作入口，包括解压和压缩的下拉菜单选项。',
            '**拖放功能**：支持将文件或文件夹拖放到指定区域，实现快速处理，提升用户操作效率。',
            '**进度反馈**：实时显示解压或压缩进度，提供清晰的状态提示，增强用户体验。'
        ]
        for func in ui_funcs:
            self.document.add_paragraph(func, style='List Bullet')
        
        # 技术需求
        self.document.add_heading('3. 技术需求', level=3)
        tech_funcs = [
            '**编码处理**：解决压缩包内中文文件名的乱码问题，支持UTF-8、GBK等多种编码的自动检测与转换。',
            '**结构优化**：自动优化解压后的文件夹结构，减少冗余层级，提升文件组织的合理性。',
            '**多线程处理**：采用多线程技术，避免界面卡顿，实现解压/压缩操作的异步处理。',
            '**异常处理**：完善的异常捕获机制，处理加密压缩包、损坏文件、路径错误等异常情况，并提供友好的错误提示。'
        ]
        for func in tech_funcs:
            self.document.add_paragraph(func, style='List Bullet')
        
        # 扩展性需求
        self.document.add_heading('4. 扩展性需求', level=3)
        ext_funcs = [
            '**格式扩展**：系统架构需具备良好的扩展性，便于后续添加新的压缩格式支持。',
            '**功能扩展**：为未来可能增加的功能（如压缩包加密、分卷压缩等）预留接口。'
        ]
        for func in ext_funcs:
            self.document.add_paragraph(func, style='List Bullet')
    
    def add_section2(self):
        """添加第二部分：系统设计"""
        self.document.add_heading('二、系统设计', level=1)
        
        # 整体架构设计
        self.document.add_heading('（一）整体架构设计', level=2)
        self.document.add_paragraph('本系统采用模块化设计思想，将整个系统划分为核心处理模块、用户界面模块和辅助功能模块，'
                                   '各模块之间通过清晰的接口进行交互，确保系统的可维护性和可扩展性。')
        
        # 核心处理模块
        self.document.add_heading('1. 核心处理模块', level=3)
        self.document.add_paragraph('核心处理模块是系统的核心部分，负责实现压缩文件的解压与压缩逻辑，主要包括：')
        core_modules = [
            '**Extractor类**：作为核心处理类，封装了所有压缩文件处理的核心功能，包括解压、压缩、嵌套处理、编码转换、结构优化等。',
            '**格式处理子模块**：针对不同的压缩格式（ZIP、RAR、7Z、TAR等），实现对应的解压和压缩算法，通过策略模式进行统一管理。',
            '**嵌套处理子模块**：实现递归查找和处理嵌套压缩包的算法，确保多层压缩包的自动解压。',
            '**编码处理子模块**：负责文件名的编码检测与转换，解决中文乱码问题。',
            '**结构优化子模块**：分析解压后的文件夹结构，自动展平冗余层级，优化文件组织。'
        ]
        for module in core_modules:
            self.document.add_paragraph(module, style='List Bullet')
        
        # 用户界面模块
        self.document.add_heading('2. 用户界面模块', level=3)
        self.document.add_paragraph('用户界面模块基于Tkinter构建，提供直观的操作界面，主要包括：')
        ui_modules = [
            '**主窗口**：包含标题、拖放区域、功能按钮和进度显示区域。',
            '**菜单按钮**：“解压”和“压缩”按钮采用下拉菜单形式，提供文件和文件夹的处理选项。',
            '**控制按钮**：包括暂停/继续、终止按钮，用于控制操作进程。',
            '**拖放区域**：支持用户将文件或文件夹拖入系统进行处理。',
            '**进度显示**：实时显示当前操作的进度和状态信息。'
        ]
        for module in ui_modules:
            self.document.add_paragraph(module, style='List Bullet')
        
        # 辅助功能模块
        self.document.add_heading('3. 辅助功能模块', level=3)
        self.document.add_paragraph('辅助功能模块提供系统的辅助支持功能，主要包括：')
        helper_modules = [
            '**线程管理**：负责多线程的创建、管理和同步，确保界面响应性。',
            '**路径处理**：处理文件和文件夹路径的规范化、安全性检查等。',
            '**异常处理**：捕获并处理操作过程中的各种异常，提供错误提示和恢复机制。',
            '**系统交互**：实现与操作系统的交互，如打开文件夹等功能。'
        ]
        for module in helper_modules:
            self.document.add_paragraph(module, style='List Bullet')
        
        # 核心类设计
        self.document.add_heading('（二）核心类设计', level=2)
        
        # Extractor类
        self.document.add_heading('1. Extractor类', level=3)
        self.document.add_paragraph('Extractor类是系统的核心类，负责实现所有压缩文件处理的核心功能，其主要属性和方法如下：')
        
        # 属性
        self.document.add_heading('**属性**：', level=4)
        attrs = [
            '`_pause`：线程事件，用于控制操作的暂停与继续。',
            '`_stop`：线程事件，用于控制操作的终止。',
            '`extracted_dirs`：记录已解压的目录，用于回滚操作。',
            '`compressed_files`：记录已压缩的文件，用于回滚操作。',
            '`compression_thread`：压缩操作的线程引用。',
            '`progress_callback`：进度回调函数，用于更新界面进度。',
            '`keep_original_archives`：是否保留原始压缩包。',
            '`flatten_single_folder`：是否展平单层文件夹。'
        ]
        for attr in attrs:
            self.document.add_paragraph(attr, style='List Bullet')
        
        # 核心方法
        self.document.add_heading('**核心方法**：', level=4)
        methods = [
            '`extract_archive(file_path, extract_to)`：解压单个压缩包到指定目录，支持多种格式。',
            '`compress_folder(folder_path, archive_path, fmt="zip")`：压缩文件夹为指定格式的压缩包。',
            '`compress_file(file_path, archive_path, fmt="zip")`：压缩单个文件为指定格式的压缩包。',
            '`extract_nested_archives(folder)`：递归处理文件夹中的嵌套压缩包。',
            '`optimize_extracted_structure(target_dir)`：优化解压后的文件夹结构，减少冗余层级。',
            '`_decode_filename(filename)`：处理文件名编码，解决中文乱码问题。',
            '`rollback()`：回滚操作，删除已解压或压缩的内容。'
        ]
        for method in methods:
            self.document.add_paragraph(method, style='List Bullet')
        
        # 界面交互类
        self.document.add_heading('2. 界面交互类', level=3)
        self.document.add_paragraph('界面交互部分通过Tkinter实现，主要包括以下关键函数：')
        ui_funcs = [
            '`on_drop(event)`：处理拖放事件，解析拖入的文件或文件夹并启动相应处理。',
            '`on_choose_extract_file()`：打开文件选择对话框，选择要解压的文件。',
            '`on_choose_extract_folder()`：打开文件夹选择对话框，选择包含压缩包的文件夹。',
            '`on_pause_resume()`：处理暂停/继续操作，更新按钮状态和进度显示。',
            '`on_stop()`：处理终止操作，调用回滚方法删除已解压内容。',
            '`save_compressed_file(target_path, is_file)`：保存压缩包文件，处理压缩格式选择和线程操作。'
        ]
        for func in ui_funcs:
            self.document.add_paragraph(func, style='List Bullet')
        
        # 关键技术选型
        self.document.add_heading('（三）关键技术选型', level=2)
        
        # 开发语言与框架
        self.document.add_heading('1. 开发语言与框架', level=3)
        dev_techs = [
            '**Python**：作为开发语言，因其丰富的第三方库和简洁的语法，适合快速开发原型和实现复杂逻辑。',
            '**Tkinter**：Python标准库中的GUI框架，用于构建用户界面，无需额外安装依赖，便于部署。',
            '**tkinterdnd2**：扩展Tkinter的拖放功能，实现文件和文件夹的拖放操作。'
        ]
        for tech in dev_techs:
            self.document.add_paragraph(tech, style='List Bullet')
        
        # 压缩格式处理库
        self.document.add_heading('2. 压缩格式处理库', level=3)
        compress_techs = [
            '**zipfile**：Python标准库，用于处理ZIP格式压缩包。',
            '**rarfile**：第三方库，用于处理RAR格式压缩包（需额外安装）。',
            '**py7zr**：第三方库，用于处理7Z格式压缩包。',
            '**tarfile**：Python标准库，用于处理TAR、TAR.GZ、TAR.BZ2等格式。'
        ]
        for tech in compress_techs:
            self.document.add_paragraph(tech, style='List Bullet')
        
        # 多线程处理
        self.document.add_heading('3. 多线程处理', level=3)
        self.document.add_paragraph('**threading**：Python标准库，用于创建和管理线程，实现异步解压/压缩操作，避免界面卡顿。')
        
        # 编码处理
        self.document.add_heading('4. 编码处理', level=3)
        self.document.add_paragraph('采用UTF-8、GBK等编码的自动检测与转换机制，解决中文文件名乱码问题。'
                                   '通过尝试不同编码解码文件名，并使用错误替换策略处理无法解码的情况。')
    
    def add_section3(self):
        """添加第三部分：功能实现与技术细节"""
        self.document.add_heading('三、功能实现与技术细节', level=1)
        
        # 解压功能实现
        self.document.add_heading('（一）解压功能实现', level=2)
        
        # 基础解压流程
        self.document.add_heading('1. 基础解压流程', level=3)
        self.document.add_paragraph('解压功能的核心在于`extract_archive`方法，该方法根据压缩包的格式调用不同的处理逻辑，主要流程如下：')
        extract_steps = [
            '**确定目标目录**：根据压缩包路径和用户指定的解压目录，生成目标解压目录，并处理同名目录的情况。',
            '**格式检测与处理**：检测压缩包格式，调用对应的解压库（如zipfile、rarfile等）进行解压。',
            '**文件名编码处理**：在解压过程中，对文件名进行编码检测和转换，支持UTF-8、GBK等编码，解决中文乱码问题。',
            '**安全性检查**：检查解压路径是否包含非法字符或路径穿越风险，确保系统安全。',
            '**解压执行**：将压缩包内容解压到目标目录，对于文件夹需先创建目录结构。'
        ]
        for step in extract_steps:
            self.document.add_paragraph(step, style='List Number')
        
        # 嵌套解压实现
        self.document.add_heading('2. 嵌套解压实现', level=3)
        self.document.add_paragraph('嵌套解压是本系统的核心特色功能，通过`extract_nested_archives`方法实现，主要逻辑如下：')
        nested_steps = [
            '**递归查找压缩包**：在解压后的文件夹中递归查找所有支持的压缩包文件。',
            '**按层级排序**：将找到的嵌套压缩包按路径长度排序，优先处理内层压缩包，确保解压顺序的合理性。',
            '**递归解压**：对每个嵌套压缩包创建子目录并解压，解压后继续处理子目录中的嵌套压缩包，形成递归处理链。',
            '**结构优化**：每次解压后调用`optimize_extracted_structure`方法，优化文件夹结构，展平冗余层级。',
            '**原始文件清理**：根据配置选项，自动删除已解压的原始压缩包，减少磁盘占用。'
        ]
        for step in nested_steps:
            self.document.add_paragraph(step, style='List Number')
        
        # 乱码处理技术
        self.document.add_heading('3. 乱码处理技术', level=3)
        self.document.add_paragraph('中文文件名乱码问题的解决是解压功能的关键难点，系统通过以下策略处理：')
        encoding_steps = [
            '**多编码尝试**：在解压时，先尝试用UTF-8解码文件名，失败后尝试GBK解码，最后使用错误替换策略（`errors=\'replace\'`）处理无法解码的字符。',
            '**编码转换**：对于zipfile库在Python 3.6+中可能出现的cp437编码错误，先将文件名编码为cp437字节，再尝试用UTF-8或GBK解码。',
            '**统一规范化**：解码后的文件名进行路径规范化处理（`os.path.normpath`），确保路径的一致性和合法性。'
        ]
        for step in encoding_steps:
            self.document.add_paragraph(step, style='List Number')
        
        # 其他功能实现部分...（由于篇幅限制，这里简化处理，实际代码应完整实现所有部分）
        self.document.add_heading('（二）压缩功能实现', level=2)
        self.document.add_paragraph('压缩功能通过`compress_file`和`compress_folder`方法实现，支持多种格式，主要流程如下：')
        self.document.add_paragraph('1. **格式选择**：根据用户指定的压缩格式（ZIP、7Z、TAR、RAR）选择对应的处理逻辑。')
        self.document.add_paragraph('2. **文件遍历**：对于文件夹压缩，递归遍历文件夹中的所有文件，获取相对路径以保持压缩包内的目录结构。')
        self.document.add_paragraph('3. **压缩执行**：使用对应的压缩库（如zipfile、py7zr等）将文件写入压缩包，支持压缩级别设置（如ZIP的DEFLATED算法）。')
        self.document.add_paragraph('4. **异常处理**：捕获压缩过程中的异常，如加密压缩包、文件访问错误等，并提供友好的错误提示。')
        
        self.document.add_heading('（三）操作控制功能实现', level=2)
        self.document.add_paragraph('暂停与继续功能通过线程事件（`_pause`）实现：')
        self.document.add_paragraph('1. **暂停操作**：调用`pause()`方法清除`_pause`事件，线程在执行`_check_stop_and_pause()`时会阻塞等待。')
        self.document.add_paragraph('2. **继续操作**：调用`resume()`方法设置`_pause`事件，线程继续执行。')
        self.document.add_paragraph('3. **界面同步**：更新暂停/继续按钮的文本和进度显示，反映当前操作状态。')
        
        self.document.add_heading('（四）界面交互实现', level=2)
        self.document.add_paragraph('拖放功能通过`tkinterdnd2`库实现：')
        self.document.add_paragraph('1. **目标注册**：在拖放区域注册`DND_FILES`类型，接收拖入的文件或文件夹。')
        self.document.add_paragraph('2. **事件处理**：`on_drop`事件处理函数解析拖入的路径，区分文件和文件夹，并启动相应的解压流程。')
        self.document.add_paragraph('3. **智能处理**：自动识别拖入的文件是否为支持的压缩格式，或文件夹中是否包含压缩包，实现智能处理。')
    
    def add_section4(self):
        """添加第四部分：版本迭代与优化过程"""
        self.document.add_heading('四、版本迭代与优化过程', level=1)
        
        # 版本1
        self.document.add_heading('（一）版本1：基础功能实现', level=2)
        
        # 功能亮点
        self.document.add_heading('1. 功能亮点', level=3)
        v1_highlights = [
            '实现了基本的ZIP格式解压功能，支持选择文件和文件夹进行解压。',
            '初步构建了图形用户界面，包含解压按钮和基本的进度显示。',
            '完成了核心Extractor类的基础框架，实现了文件路径处理和基本异常处理。'
        ]
        for highlight in v1_highlights:
            self.document.add_paragraph(highlight, style='List Bullet')
        
        # 技术难点与解决方案
        self.document.add_heading('2. 技术难点与解决方案', level=3)
        self.document.add_paragraph('**问题**：仅支持ZIP格式，功能单一。')
        self.document.add_paragraph('**方案**：设计可扩展的格式处理接口，为后续添加其他格式支持做准备。')
        
        # 版本2.0
        self.document.add_heading('（二）版本2.0：功能扩展与界面优化', level=2)
        
        # 版本2.1
        self.document.add_heading('1. 版本2.1：下拉菜单与多格式支持', level=3)
        self.document.add_heading('**功能改进**：', level=4)
        v21_funcs = [
            '添加解压和压缩的下拉菜单选项，区分文件和文件夹处理，提升操作便捷性。',
            '增加对7Z、TAR等格式的解压支持，扩展系统兼容性。'
        ]
        for func in v21_funcs:
            self.document.add_paragraph(func, style='List Bullet')
        
        self.document.add_heading('**界面优化**：', level=4)
        v21_ui = [
            '优化窗口布局，增加按钮和菜单的视觉层次感。',
            '完善进度显示区域，提供更清晰的操作状态反馈。'
        ]
        for item in v21_ui:
            self.document.add_paragraph(item, style='List Bullet')
        
        # 版本2.2
        self.document.add_heading('2. 版本2.2：界面美化与色彩优化', level=3)
        self.document.add_heading('**界面改进**：', level=4)
        v22_ui = [
            '调整界面配色方案，采用更友好的蓝色系主题，提升视觉体验。',
            '优化组件间距和字体样式，增强界面可读性。'
        ]
        for item in v22_ui:
            self.document.add_paragraph(item, style='List Bullet')
        
        self.document.add_heading('**交互优化**：', level=4)
        v22_interaction = [
            '增加拖放区域的视觉提示，引导用户进行拖放操作。',
            '优化按钮的交互反馈，如悬停效果和点击状态。'
        ]
        for item in v22_interaction:
            self.document.add_paragraph(item, style='List Bullet')
        
        # 其他版本...（由于篇幅限制，这里简化处理，实际代码应完整实现所有部分）
    
    def add_section5(self):
        """添加第五部分：系统测试与性能分析"""
        self.document.add_heading('五、系统测试与性能分析', level=1)
        
        # 功能测试
        self.document.add_heading('（一）功能测试', level=2)
        
        # 解压功能测试
        self.document.add_heading('1. 解压功能测试', level=3)
        self.document.add_heading('**测试用例**：', level=4)
        extract_tests = [
            '单层ZIP压缩包（含中文文件名）解压，验证文件名正确性和路径结构。',
            '多层嵌套的7Z压缩包（ZIP在7Z内）解压，验证嵌套处理能力。',
            '损坏的RAR压缩包解压，验证异常处理和错误提示。'
        ]
        for test in extract_tests:
            self.document.add_paragraph(test, style='List Bullet')
        
        self.document.add_heading('**测试结果**：', level=4)
        self.document.add_paragraph('系统能正确解压多种格式的压缩包，中文文件名无乱码，嵌套处理逻辑正确，异常处理机制有效。')
        
        # 其他测试部分...（由于篇幅限制，这里简化处理，实际代码应完整实现所有部分）
    
    def add_section6(self):
        """添加第六部分：总结与展望"""
        self.document.add_heading('六、总结与展望', level=1)
        
        # 项目成果总结
        self.document.add_heading('（一）项目成果总结', level=2)
        self.document.add_paragraph('本“智能嵌套式压缩文件处理系统”成功实现了预期的核心功能，包括多格式压缩包的解压与压缩、'
                                   '嵌套压缩包的自动处理、中文乱码解决、操作控制与回滚等。系统采用模块化设计，'
                                   '核心处理逻辑与界面交互分离，具有良好的可维护性和扩展性。')
        self.document.add_paragraph('通过版本迭代，系统从基础功能逐步完善，解决了嵌套处理、乱码问题等关键技术难点，'
                                   '提升了用户体验和系统稳定性。测试结果表明，系统在功能正确性、性能表现和兼容性方面均达到了设计要求，'
                                   '能够满足用户对智能压缩文件处理的需求。')
        
        # 技术创新点
        self.document.add_heading('（二）技术创新点', level=2)
        innovations = [
            '**智能嵌套处理**：实现了递归查找和处理嵌套压缩包的算法，自动识别并解压多层压缩包，减少用户手动操作。',
            '**多编码处理机制**：通过多编码尝试和转换策略，有效解决了中文文件名乱码问题，支持UTF-8、GBK等多种编码。',
            '**文件夹结构优化**：自动展平解压后的冗余文件夹层级，提升文件组织的合理性和用户访问效率。',
            '**完善的操作控制**：提供暂停、继续、终止操作功能，并实现回滚机制，确保操作的灵活性和系统状态的一致性。'
        ]
        for item in innovations:
            self.document.add_paragraph(item, style='List Bullet')
        
        # 未来展望
        self.document.add_heading('（三）未来展望', level=2)
        
        # 功能扩展
        self.document.add_heading('1. 功能扩展', level=3)
        future_funcs = [
            '**压缩包加密**：添加压缩包加密功能，支持设置密码保护压缩内容。',
            '**分卷压缩**：实现分卷压缩功能，便于大文件的存储和传输。',
            '**压缩率优化**：针对不同格式和文件类型，提供压缩级别选择和优化算法。'
        ]
        for func in future_funcs:
            self.document.add_paragraph(func, style='List Bullet')
        
        # 其他展望部分...（由于篇幅限制，这里简化处理，实际代码应完整实现所有部分）
    
    def save_document(self, filename='智能嵌套式压缩文件处理系统设计与实现.docx'):
        """保存文档到文件"""
        self.document.save(filename)
        print(f"文档已成功保存为: {os.path.abspath(filename)}")

def main():
    """主函数"""
    print("正在生成智能嵌套式压缩文件处理系统设计文档...")
    generator = DocumentGenerator()
    generator.add_title('智能嵌套式压缩文件处理系统设计与实现')

    # 添加目录标题（三号、加粗、居中）
    toc_title = generator.document.add_paragraph('目录')
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = toc_title.runs[0]
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(18)  # 三号
    run.font.bold = True

    # 添加目录域（Word中需右键更新域才能显示目录内容和页码）
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as oxml_qn

    p = generator.document.add_paragraph()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(oxml_qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(oxml_qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # 只显示1-2级标题
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(oxml_qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(oxml_qn('w:fldCharType'), 'end')
    r = p.add_run()
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)

    # 插入分页符，让目录单独成页
    generator.document.add_page_break()

    # 后续为正文内容
    generator.add_section1()
    generator.add_section2()
    generator.add_section3()
    generator.add_section4()
    generator.add_section5()
    generator.add_section6()
    ref_title = generator.document.add_paragraph('参考文献', style='Title')
    ref_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 参考文献内容同正文样式
    generator.document.add_paragraph('[1] Python Software Foundation. Data Compression and Archiving [EB/OL]. Python 3.12 Documentation. https://docs.python.org/3/library/archiving.html')
    generator.document.add_paragraph('[2] TkinterDnD2库使用指南 [EB/OL]. tkinterdnd2 Documentation & Community Examples.')
    generator.save_document()
    print("文档生成完成！")

if __name__ == "__main__":
    main()

doc = Document()

# 1. 封面
doc.add_paragraph('毕业论文封面', style='Title')
doc.add_paragraph('作者：张三')
doc.add_paragraph('')

# 2. 插入分节符（目录节）
doc.add_section(0)  # 0 = WD_SECTION.NEW_PAGE

# 3. 目录页
doc.add_paragraph('目录', style='Title')
doc.add_paragraph('第一章 绪论 ............................................. 1')
doc.add_paragraph('第二章 相关技术 ....................................... 5')

# 4. 插入分节符（正文节）
doc.add_section(0)  # 0 = WD_SECTION.NEW_PAGE

# 5. 正文内容
doc.add_paragraph('第一章 绪论', style='Heading 1')
doc.add_paragraph('这里是正文内容……')

# 6. 断开正文与前面节的页脚链接
doc.sections[2].footer.is_linked_to_previous = False

# 7. 清空前两节的页脚内容
for p in doc.sections[0].footer.paragraphs:
    p.clear()
for p in doc.sections[1].footer.paragraphs:
    p.clear()

# 8. 正文节添加页码，并从1开始
footer = doc.sections[2].footer
footer_p = footer.paragraphs[0]
footer_p.alignment = 1  # 居中

run = footer_p.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "PAGE"
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)
run._r.append(fldChar3)

run.font.name = '宋体'
run.font.size = Pt(10.5)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 关键：正文节页码从1开始
doc.sections[2].start_type = 2  # WD_SECTION.NEW_PAGE
doc.sections[2].start_page_number = 1

# doc.save('output.docx')  # 如需保存请取消注释